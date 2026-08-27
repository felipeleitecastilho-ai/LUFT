from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# CAPA
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('KEYRUS AI x LUFT LOGISTICS')
run.font.size = Pt(20)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Licoes Aprendidas')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(236, 40, 73)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Projeto Data Lakehouse Luft Logistics')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('Documento interno Keyrus | Agosto 2026').font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# SUMARIO
doc.add_heading('Sumario', level=1)
for item in ['1. Cortex Agent - Como criar', '2. ETL - Regras de ouro', '3. PowerShell no CoCo', '4. Como me comunicar com o CoCo', '5. Organizacao de projeto', '6. Setup inicial', '7. Nomes de tabelas e colunas', '8. Monitoramento', '9. Checklist novo projeto']:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# SECAO 1
doc.add_heading('1. Cortex Agent - Como criar corretamente', level=1)
doc.add_heading('Sintaxe que funciona:', level=2)
doc.add_paragraph('Use sempre FROM SPECIFICATION com YAML (nunca JSON)', style='List Bullet')
doc.add_paragraph('CREATE OR REPLACE AGENT ... FROM SPECIFICATION $$ YAML $$', style='List Bullet')
doc.add_paragraph('Verificar com DESCRIBE AGENT se agent_spec nao esta vazio', style='List Bullet')
doc.add_heading('O que NAO funciona:', level=2)
doc.add_paragraph("spec = '...' - nao persiste", style='List Bullet')
doc.add_paragraph('SPECIFICATION = $$ JSON $$ - nao salva', style='List Bullet')
doc.add_heading('Evitar alucinacao:', level=2)
doc.add_paragraph('NUNCA ter duas ferramentas com dados parecidos', style='List Bullet')
doc.add_paragraph('Se trocar fonte, SUBSTITUIR (nao adicionar segunda)', style='List Bullet')
doc.add_paragraph('Roteamento claro nas instructions', style='List Bullet')
doc.add_heading('Grants:', level=2)
doc.add_paragraph('CREATE OR REPLACE apaga todos os grants!', style='List Bullet')
doc.add_paragraph('Sempre re-aplicar GRANT USAGE apos recriar', style='List Bullet')
doc.add_paragraph('Agente usa DEFAULT_ROLE do usuario (nao a role da sessao)', style='List Bullet')
doc.add_page_break()

# SECAO 2
doc.add_heading('2. ETL - Regras de ouro', level=1)
doc.add_heading('Volume:', level=2)
doc.add_paragraph('Ate 50k: puxa tudo de uma vez', style='List Bullet')
doc.add_paragraph('50k-500k: lotes (fetchmany)', style='List Bullet')
doc.add_paragraph('500k-1M: OFFSET/FETCH', style='List Bullet')
doc.add_paragraph('Mais de 1M: filtrar por data (7 ou 90 dias)', style='List Bullet')
doc.add_heading('Modos de carga:', level=2)
doc.add_paragraph('Full: apaga e recarrega', style='List Bullet')
doc.add_paragraph('Incremental: so dados novos', style='List Bullet')
doc.add_paragraph('Janela N dias: ultimos N dias (GPS, telemetria)', style='List Bullet')
doc.add_heading('ANTES de extrair, SEMPRE:', level=2)
doc.add_paragraph('Contar registros (COUNT)', style='List Bullet')
doc.add_paragraph('Verificar formato de datas', style='List Bullet')
doc.add_paragraph('Verificar se valores sao texto ou numero', style='List Bullet')
doc.add_page_break()

# SECAO 3
doc.add_heading('3. PowerShell no CoCo', level=1)
doc.add_paragraph('NAO funciona: heredoc, &&, aspas mistas', style='List Bullet')
doc.add_paragraph('Funciona: Set-Content, ; (ponto e virgula), aspas simples no git', style='List Bullet')
doc.add_page_break()

# SECAO 4
doc.add_heading('4. Como me comunicar com o CoCo', level=1)
doc.add_heading('Pedir ANTES de implementar:', level=2)
doc.add_paragraph('"Me diz o que vai fazer antes de fazer"', style='List Bullet')
doc.add_paragraph('"Faz backup antes de alterar"', style='List Bullet')
doc.add_paragraph('"Testa antes de dizer que esta pronto"', style='List Bullet')
doc.add_paragraph('"Nao faz commit sem perguntar"', style='List Bullet')
doc.add_paragraph('"Verifica volume e formato antes de assumir"', style='List Bullet')
doc.add_heading('Fluxo nova fonte:', level=2)
doc.add_paragraph('1. Listar tabelas > 2. Eu escolho > 3. Verifica volume > 4. Define modo > 5. Implementa > 6. Testa > 7. Entrega')
doc.add_page_break()

# SECAO 5
doc.add_heading('5. Organizacao de projeto', level=1)
doc.add_paragraph('scripts_etl/ - Scripts Python e .bat', style='List Bullet')
doc.add_paragraph('docs/ - Documentacao', style='List Bullet')
doc.add_paragraph('streamlit/ - Apps Streamlit', style='List Bullet')
doc.add_paragraph('dados_csv/ - Arquivos CSV', style='List Bullet')
doc.add_paragraph('backups/ - Backups SQL', style='List Bullet')
doc.add_page_break()

# SECAO 6
doc.add_heading('6. Setup inicial (fazer PRIMEIRO)', level=1)
doc.add_paragraph('1. Criar pasta no Git')
doc.add_paragraph('2. Criar .gitignore (credenciais, csv, temp)')
doc.add_paragraph('3. Criar estrutura de pastas')
doc.add_paragraph('4. Commit inicial + push')
doc.add_heading('Seguranca:', level=2)
doc.add_paragraph('BACKUP antes de DROP/ALTER/CREATE OR REPLACE', style='List Bullet')
doc.add_paragraph('PERGUNTAR antes de commit e push', style='List Bullet')
doc.add_paragraph('NUNCA commitar credenciais', style='List Bullet')
doc.add_page_break()

# SECAO 7
doc.add_heading('7. Nomes de tabelas e colunas', level=1)
doc.add_paragraph('Bronze: NOME_FONTE_RAW', style='List Bullet')
doc.add_paragraph('Silver: NOME (sem sufixo)', style='List Bullet')
doc.add_paragraph('Gold: VW_SEMANTICA_NOME, SEMANTICO_NOME', style='List Bullet')
doc.add_paragraph()
doc.add_paragraph('Prefixos: CD_=codigo, NM_=nome, VL_=valor, QT_=quantidade, DT_=data, FL_=flag')
doc.add_paragraph('Na Gold: remover prefixos (NM_FILIAL vira FILIAL)')
doc.add_page_break()

# SECAO 8
doc.add_heading('8. Monitoramento (criar SEMPRE)', level=1)
doc.add_paragraph('ETL_CONTROL: ultima data de carga por fonte', style='List Bullet')
doc.add_paragraph('ETL_LOG: historico de cada execucao', style='List Bullet')
doc.add_paragraph('Alerta email 07h (resumo) e 08h (falhas)', style='List Bullet')
doc.add_paragraph('Painel Streamlit com status visual', style='List Bullet')
doc.add_page_break()

# SECAO 9
doc.add_heading('9. Checklist novo projeto', level=1)
for item in ['Definir fontes e volumes', 'Criar database + schemas', 'Criar roles + future grants', 'Key Pair para ETL', 'ETL Bronze', 'Dynamic Tables Silver', 'Views Gold + JOINs', 'Semantic Views', 'Criar Agente (YAML)', 'Testar agente', 'Alertas email', 'Painel Streamlit', 'Agendar ETL', 'Documentar + KT']:
    doc.add_paragraph(item, style='List Bullet')

out = r'C:\Users\Felipe.Santos\OneDrive - Keyrus\Área de Trabalho\PROJETOS\LUFT\docs\Licoes_Aprendidas_Keyrus.docx'
doc.save(out)
print(f'Salvo: {out}')
